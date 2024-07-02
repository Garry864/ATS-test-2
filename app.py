import streamlit as st
from streamlit_option_menu import option_menu
import google.generativeai as genai
import os
import base64
import PyPDF2 as pdf
from docx import Document
from dotenv import load_dotenv

load_dotenv()  # Load all our environment variables

api_key = os.getenv("GOOGLE_API_KEY")
if api_key:
    genai.configure(api_key=api_key)
else:
    st.error("Google API key is missing! Please check your .env file.")

# Settings
page_title = "Smart ATS"
page_icon = ":pager:"
page2 = ":desktop_computer:"
layout = "centered"

st.set_page_config(page_title=page_title, page_icon=page_icon, layout=layout)
st.title(f"{page_title} {page2}")
st.markdown("Improve Your Resume with ATS")

# Initialize session state variables
if 'resume_analysis_response' not in st.session_state:
    st.session_state.resume_analysis_response = None

if 'percentage_match_response' not in st.session_state:
    st.session_state.percentage_match_response = None

if 'improvisation_response' not in st.session_state:
    st.session_state.improvisation_response = None

# Sidebar
st.sidebar.title("Upload Your Resume")
uploaded_file = st.sidebar.file_uploader("Upload File", type=["pdf", "docx"], help="Please upload your resume in PDF or DOCX format")

# Gemini Setup
def get_gemini_response(input):
    try:
        model = genai.GenerativeModel('gemini-pro')
        response = model.generate_content(input)
        return response.text
    except Exception as e:
        st.error(f"An error occurred while generating response: {e}")
        return None

def input_pdf_text(uploaded_file):
    try:
        reader = pdf.PdfReader(uploaded_file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        if not text.strip():
            st.error("The PDF file does not contain extractable text.")
        return text
    except Exception as e:
        st.error(f"Error reading PDF file: {e}")
        return ""

def input_docx_text(uploaded_file):
    try:
        doc = Document(uploaded_file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        st.error(f"Error reading DOCX file: {e}")
        return ""

def show_pdf(uploaded_file):
    try:
        pdf_contents = uploaded_file.read()
        base64_pdf = base64.b64encode(pdf_contents).decode('utf-8')
        pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="700" height="800" type="application/pdf"></iframe>'
        st.markdown(pdf_display, unsafe_allow_html=True)
    except Exception as e:
        st.error(f"Error displaying PDF file: {e}")

def show_docx(uploaded_file):
    try:
        doc = Document(uploaded_file)
        for paragraph in doc.paragraphs:
            st.write(paragraph.text)
    except Exception as e:
        st.error(f"Error displaying DOCX file: {e}")

# Prompt Templates
input_prompt1 = """
You are an experienced Technical Human Resource Manager. Your task is to review the provided resume against the job description. 
Please share your professional evaluation on whether the candidate's profile aligns with the role. 
Highlight the strengths and weaknesses of the applicant in relation to the specified job requirements. Start with the name and job profile of the candidate
resume:{text}
description:{jd}
"""

input_prompt2 = """
You are a skilled ATS (Applicant Tracking System) scanner with a deep understanding of data science and ATS functionality. 
Your task is to evaluate the resume against the provided job description. Give me the percentage of match if the resume matches
the job description. First, the output should come as a percentage, then keywords missing, and finally, final thoughts.
resume:{text}
description:{jd}
"""

input_prompt3 = """
Consider yourself as an ATS advisor for the application. Give the best possible improvisation suggestions from the response and form the resume and description. Highlight the skillsets which need to be improved according to the description.
resume:{text}
description:{jd}
response:{response}
"""

# Layout
input_text = st.text_area("Job Description:", key="input", height=200)

# Option Menu for Analysis Types
selected = option_menu(
    menu_title=None,
    options=["Resume Analysis", "Percentage Match", "Improvisation"],
    icons=["file-earmark-pdf", "percent", "lightbulb"],
    orientation="horizontal"
)

# Perform actions based on selected menu item
if uploaded_file is not None:
    pdf_content = ""
    if uploaded_file.type == "pdf":
        pdf_content = input_pdf_text(uploaded_file)
    elif uploaded_file.type == "docx":
        pdf_content = input_docx_text(uploaded_file)

    if selected == "Resume Analysis":
        show_pdf(uploaded_file)
        with st.expander("About this resume ⤵"):
            input_data = {
                "text": pdf_content,
                "jd": input_text
            }
            response = get_gemini_response(input_prompt1.format(**input_data))
            if response:
                st.session_state.resume_analysis_response = response
                st.subheader("Resume Analysis")
                st.write(response)

    elif selected == "Percentage Match":
        show_pdf(uploaded_file)
        with st.expander("ATS Score of this Resume ⤵"):
            input_data = {
                "text": pdf_content,
                "jd": input_text
            }
            response = get_gemini_response(input_prompt2.format(**input_data))
            if response:
                st.session_state.percentage_match_response = response
                st.subheader("Percentage Match")
                st.write(response)

    elif selected == "Improvisation":
        if st.session_state.percentage_match_response:
            show_pdf(uploaded_file)
            with st.expander("Suggestions for your Improvisation here ⤵"):
                input_data = {
                    "text": pdf_content,
                    "jd": input_text,
                    "response": st.session_state.percentage_match_response
                }
                response = get_gemini_response(input_prompt3.format(**input_data))
                if response:
                    st.subheader("Improvisation Suggestions")
                    st.write(response)
                    st.session_state.improvisation_response = response
        else:
            st.warning("Please complete the 'Percentage Match' analysis first.")

else:
    st.warning("Please upload the resume")

# Download links
if st.session_state.resume_analysis_response:
    st.sidebar.download_button(
        label="Download Resume Analysis",
        data=st.session_state.resume_analysis_response,
        file_name="resume_analysis.txt"
    )

if st.session_state.percentage_match_response:
    st.sidebar.download_button(
        label="Download Percentage Match",
        data=st.session_state.percentage_match_response,
        file_name="percentage_match.txt"
    )

if st.session_state.improvisation_response:
    st.sidebar.download_button(
        label="Download Improvisation Suggestions",
        data=st.session_state.improvisation_response,
        file_name="improvisation_suggestions.txt"
    )

# Footer
st.sidebar.markdown("---")
st.sidebar.text("Built with ❤️ by Gaurav Yadav")
